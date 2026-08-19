import io
from docx import Document as DocxDocument
from app.services.extractors.base import BaseDocumentExtractor, ExtractedDocument, ExtractedPage


class DOCXExtractor(BaseDocumentExtractor):
    """DOCX Document Extractor."""

    def extract(self, file_bytes: bytes) -> ExtractedDocument:
        if not file_bytes:
            raise ValueError("File content is empty")

        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            
            if not full_text:
                raise ValueError("DOCX contains no readable text")

            # DOCX files do not have explicit page breaks, single page container
            pages = [ExtractedPage(page_number=1, text=full_text)]
            return ExtractedDocument(page_count=1, pages=pages)
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")
